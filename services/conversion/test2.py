import io
import os

import cv2
import fitz
import numpy as np
import pdfplumber
import zlib
from PIL import Image
from docx.shared import Cm
import logging

logging.basicConfig(
    filename="app.log",  # Имя файла для логов
    filemode="a",        # Режим записи: "a" для добавления, "w" для перезаписи
    format="%(asctime)s - %(levelname)s - %(message)s",  # Формат записи
    level=logging.INFO   # Уровень логирования (DEBUG, INFO, WARNING, ERROR, CRITICAL)
)


def configure_page_from_pdf(doc):
    """
    Настройка страниц Word-документа для соответствия размеру PDF.
    :param doc: Word-документ
    """
    section = doc.sections[0]
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)

def get_pdf_page_size(pdf_path, page_number):
    """
    Возвращает размеры страницы PDF в точках.
    :param pdf_path: Путь к PDF-файлу
    :param page_number: Номер страницы (начиная с 1)
    :return: (width, height) в точках
    """
    doc = fitz.open(pdf_path)
    page = doc[page_number - 1]
    width, height = page.rect.width, page.rect.height
    doc.close()
    return width, height

def insert_image(doc, image_bytes, img_width, img_height, left, top):
    """
    Вставляет изображение в документ Word с использованием стандартного позиционирования.
    :param doc: Word-документ
    :param image_bytes: Байты изображения
    :param img_width: Ширина изображения в см
    :param img_height: Высота изображения в см
    :param left: Горизонтальная позиция в см
    :param top: Вертикальная позиция в см
    """
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(image_bytes, width=Cm(img_width), height=Cm(img_height))

def recoverpix(doc, item):
    xref = item[0]  # xref of PDF image
    smask = item[1]  # xref of its /SMask
    # special case: /SMask or /Mask exists
    if smask > 0:
        pix0 = fitz.Pixmap(doc.extract_image(xref)["image"])

        if pix0.alpha:  # catch irregular situation
            pix0 = fitz.Pixmap(pix0, 0)  # remove alpha channel
        mask = fitz.Pixmap(doc.extract_image(smask)["image"])

        try:
            pix = fitz.Pixmap(pix0, mask)
        except:  # fallback to original base image in case of problems
            pix = fitz.Pixmap(doc.extract_image(xref)["image"])

        if pix0.n > 3:
            ext = "pam"
        else:
            ext = "png"

        return {  # create dictionary expected by caller
            "ext": ext,
            "colorspace": pix.colorspace.n,
            "image": pix.tobytes(ext),
        }

    # special case: /ColorSpace definition exists
    # to be sure, we convert these cases to RGB PNG images
    if "/ColorSpace" in doc.xref_object(xref, compressed=True):
        pix = fitz.Pixmap(doc, xref)
        pix = fitz.Pixmap(fitz.csRGB, pix)
        return {  # create dictionary expected by caller
            "ext": "png",
            "colorspace": 3,
            "image": pix.tobytes("png"),
        }
    return doc.extract_image(xref)

def has_same_image(image_file_dict, img_data):
    for image_path, image_info in image_file_dict.items():
        with open(image_path, "rb") as f1:
            if f1.read() == img_data:
                return image_path
    return None

def extract_images(doc, imgdir="images"):
    dimlimit = 0  # 100  # each image side must be greater than this
    relsize = 0  # 0.05  # image : image size ratio must be larger than this (5%)
    abssize = 0  # 2048  # absolute image size limit 2 KB: ignore if smaller
    page_count = doc.page_count
    xreflist = []
    result = []
    image_file_list = {}
    image_index = 1
    for pno in range(page_count):
        imglist = []
        il = doc.get_page_images(pno)
        for img in il:
            xref = img[0]
            if xref in xreflist:
                continue
            width = img[2]
            height = img[3]
            if min(width, height) <= dimlimit:
                continue
            image = recoverpix(doc, img)
            n = image["colorspace"]
            imgdata = image["image"]

            if len(imgdata) <= abssize:
                continue
            if len(imgdata) / (width * height * n) <= relsize:
                continue

            same_image = has_same_image(image_file_list, imgdata)

            if same_image is None:
                imgfile = os.path.join(imgdir, "img%05i.%s" % (image_index, image["ext"]))
                with open(imgfile, "wb") as fout:
                    fout.write(imgdata)
                image["image"] = imgfile
                image_file_list[imgfile] = image
                image_index += 1
            else:
                image = image_file_list[same_image]

            imglist.append(image)

        result.append(imglist)
        break
    return result

def extract_image(img):
    xref = img[0]
    base_image = doc.extract_image(xref)
    try:
        image_bytes = base_image["image"]
        print(f"Изображение: {base_image.get('ext', 'png')}, Цветовое пространство: {base_image.get('colorspace')}")
        image = Image.open(io.BytesIO(image_bytes))

        # Работа с smask (мягкая маска прозрачности)
        if "smask" in base_image:
            smask_xref = base_image["smask"]
            smask_image = doc.extract_image(smask_xref)
            smask_bytes = smask_image["image"]
            mask = Image.open(io.BytesIO(smask_bytes))
            image.putalpha(mask)

        # Работа с mask (жёсткая маска прозрачности)
        if "mask" in base_image:
            mask_xref = base_image["mask"]
            mask_image = doc.extract_image(mask_xref)
            mask_bytes = mask_image["image"]
            mask = Image.open(io.BytesIO(mask_bytes))
            image.putalpha(mask)

    except Exception as e:
        logging.error(f"Ошибка при обработке изображения: {e}")
        return None

    return image

def decompress_data(data):
    try:
        return zlib.decompress(data)
    except Exception as e:
        print(f"Ошибка декомпрессии: {e}")
        return None

def get_image(image_data):
    width, height, colorspace = image_data.attrs["Width"], image_data.attrs["Height"], image_data.attrs["ColorSpace"]
    mode = None
    icc_profile = None
    img = None

    if isinstance(colorspace, list):
        if colorspace[0].name == "ICCBased" and len(colorspace) > 1:
            icc_ref = colorspace[1]
            if hasattr(icc_ref, "resolve"):
                icc_profile = icc_ref.resolve().get_data()
        colorspace = colorspace[0].name
    elif hasattr(colorspace, "name"):
        colorspace = colorspace.name

    match colorspace:
        case "ICCBased" | "DeviceRGB":
            mode = "RGB"
        case "DeviceGray":
            mode = "L"
        case "DeviceCMYK":
            mode = "CMYK"

    if mode:
        try:
            img = Image.frombytes(mode, (width, height), image_data.get_data())
        except Exception as e:
            print(f"Ошибка при создании изображения: {e}")

    if "SMask" in image_data:
        mask = get_image(image_data['SMask'].resolve())
        img.putalpha(mask)

    return img

def extract_image2(img):
    image_data = img["stream"]
    img = get_image(image_data)

    img.save('imgpdf.png', 'png')

    print(img)


folder_path = "tmp/s.pdf"
output_path = "tmp/s.docx"
output_image_path = "tmp/extracted_image.png"

doc = fitz.open(folder_path)
page = doc[0]
images = extract_images(doc)
page_number = 0
#for page_number in range(len(doc)):
page = doc[page_number]
blocks = page.get_text("dict")["blocks"]
drawings = page.get_drawings()
c = page.get_contents()
print(c)
for drawing in drawings:
    for item in drawing["items"]:
        print(item)
        if item[0] == "image":  # Это изображение
            xref = item[1]  # XREF изображения
            bbox = item[2]  # Координаты экземпляра (x0, y0, x1, y1)
for block in blocks:
    print(block)

with pdfplumber.open(folder_path) as pdf:
    first_page = pdf.pages[0]
    objects = first_page.objects
    for image in objects.get("image", []):
        print(image['stream'])
        img = extract_image2(image)
        print(img)
    #for object in objects:
        #print(objects[object])
        #print(img['stream']['SMask'].resolve().get_data())
        #image = extract_image2(img)
        #image.save(f"final_image_with_transparency{0}_{img_index}.png")






