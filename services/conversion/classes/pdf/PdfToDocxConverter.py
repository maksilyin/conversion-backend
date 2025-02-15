import os
import math

import numpy as np
import pdfplumber
from docx import Document
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml import register_element_cls, OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from classes.docx.CT_Anchor import CT_Anchor
from classes.docx.CT_Line import CT_Line
from classes.docx.CT_Rect import CT_Rect
from classes.pdf.PdfParser import PdfParser


class PdfToDocxConverter:
    def __init__(self, pdf_path, output_path, image_dir):
        self.__saved_images = []
        self.__pdf_path = pdf_path
        self.__output_path = output_path
        self.__image_dir = image_dir
        self.__parser = PdfParser()
        register_element_cls('wp:anchor', CT_Anchor)
        self.__doc = Document()

        self._empty_p_size = 14
        self._empty_line_spacing = 1
        self._line_spacing_multiplier = 1.25
        self._empty_p_height = self._empty_p_size
        self.__empty_p_count = 0
        self.__empty_p = None

    def configure_page_from_pdf(self, page):
        if len(self.__doc.sections) < page.page_number:
            self.__doc.add_section()
        #section = self.__doc.sections[0]
        section = self.__doc.sections[page.page_number - 1]
        width = self.points_to_emu(page.width)
        height = self.points_to_emu(page.height)

        section.page_width = width
        section.page_height = height

        section.left_margin = 0
        section.right_margin = 0
        section.top_margin = 0
        section.bottom_margin = 0

    def get_pos_data(self, block):
        x_cm = self.points_to_emu(block['x0'])
        y_cm = self.points_to_emu(block['top'])
        height = self.points_to_emu(block['height'])
        width = self.points_to_emu(block['width'])

        return x_cm, y_cm, width, height

    def has_saved_image(self, raw_data):
        for image_data in self.__saved_images:
            if image_data['raw_data'] == raw_data:
                return image_data['path']
        return None

    def extract_image(self, img):
        image_data = img["stream"]
        raw_data = image_data.get_data()
        same_image = self.has_saved_image(raw_data)

        if same_image:
            return same_image

        img = self.__parser.get_image(image_data)
        img_format = img.format
        if img_format is None:
            img_format = 'png'
        img_path = os.path.join(self.__image_dir, "img%04i.%s" % (len(self.__saved_images) + 1, img_format.lower()))

        self.__saved_images.append({
            'path': img_path,
            'raw_data': raw_data
        })
        os.makedirs(self.__image_dir, exist_ok=True)
        img.save(img_path, img.format)
        return img_path

    def points_to_emu(self, points):
        return int(math.floor(points) * 12700)

    def points_to_cm(self, points):
        return points * 2.54 / 72

    def new_pic_anchor(self, part, image_descriptor, width, height, pos_x, pos_y):
        rId, image = part.get_or_add_image(image_descriptor)
        cx, cy = image.scaled_dimensions(width, height)
        shape_id, filename = part.next_id, image.filename
        return CT_Anchor.new_pic_anchor(shape_id, rId, filename, cx, cy, pos_x, pos_y)

    def add_float_picture(self, p, image_path_or_stream, width=None, height=None, pos_x=0, pos_y=0):
        run = p.add_run()
        anchor = self.new_pic_anchor(run.part, image_path_or_stream, width, height, pos_x, pos_y)
        run._r.add_drawing(anchor)

    def convert_pdf_to_docx_coords(self, y_pdf, page_height_pdf, unit_conversion_factor=1):
        return (page_height_pdf - y_pdf) * unit_conversion_factor

    def add_empty_paragraph(self, p_height=None, p_size=None, is_exact=False, box=None):
        if p_height is None:
            p_height = self._empty_p_size
        if p_size is None:
            p_size = self._empty_p_size

        p = self.__doc.add_paragraph()
        paragraph_format = p.paragraph_format
        if is_exact:
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            paragraph_format.line_spacing = Pt(p_height)
        else:
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
            paragraph_format.line_spacing = Pt(p_size)

        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        run = p.add_run(" ")
        font = run.font
        font.name = 'Arial'
        font.size = Pt(p_size)

        if box is not None:
            box.append(p._element)
        return p

    def get_empty_paragraph(self):
        if self.__empty_p is None:
            self.__empty_p = self.add_empty_paragraph(0.1, 0.1, True)
        return self.__empty_p

    def add_image_block(self, block, img_path, page):
        x_cm = self.points_to_emu(block['x0'])
        y_cm = self.points_to_emu(block['top'])
        height = self.points_to_emu(block['height'])
        width = self.points_to_emu(block['width'])
        p = self.get_empty_paragraph()
        self.add_float_picture(p, img_path, width, height, x_cm, y_cm)

    def is_within_margin(self, value, target, margin):
        return abs(value - target) <= margin

    def set_offset(self, top, line_height=None, box=None):
        if line_height is None:
            line_height = self._empty_p_height
        empty_paragraphs_count = int(top // line_height)
        space_before = Pt(top % line_height)

        for _ in range(empty_paragraphs_count):
            self.add_empty_paragraph(box=box)

        return space_before


    def add_paragraph(self, paragraph, page, prev_paragraph=None, box=None, set_offset=True):
        space_before = Pt(0)
        #print(paragraph)
        if set_offset:
            if prev_paragraph is None:
                space_before = self.set_offset(paragraph['top'], None, box)
            else:
                gap = paragraph['top'] - prev_paragraph['bottom']
                if gap >= prev_paragraph['size']:
                    line_height = prev_paragraph['size'] * self._line_spacing_multiplier
                    space_before = self.set_offset(gap, line_height)
                else:
                    space_before = Pt(gap)

        p = self.__doc.add_paragraph()

        for run_data in paragraph['runs']:
            #print(run_data['text'])
            run = p.add_run(run_data['text'])
            fontname = run_data['fontname'].split('-')[0].split(',')[0]

            if run_data['hyperlink'] is not None:
                run = self.add_hyperlink(p, run, run_data['hyperlink'])

            run.font.name = fontname
            run.font.size = Pt(run_data['fontsize'])
            run.underline = run_data['underline']
            run.bold = run_data['bold']

            if "Italic" in run_data['fontname'] or "Oblique" in run_data['fontname']:
                run.italic = True

            if run_data['color']:
                try:
                    rgb_color = tuple(int(c * 255) for c in run_data['color'])
                    if len(rgb_color) == 3:
                        run.font.color.rgb = RGBColor(*rgb_color)
                except (ValueError, TypeError) as e:
                    print(f"Ошибка при обработке цвета: {e}")
                    print(f"Данные цвета: {run_data['color']}")

            if run_data['letter_spacing'] is not None and not np.isnan(run_data['letter_spacing']):
                run = self.set_letter_spacing(run, run_data['letter_spacing'])

        p_format = p.paragraph_format
        p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p_format.line_spacing = Pt(paragraph['size'])
        p_format.space_after = Pt(0)
        p_format.space_before = space_before
        offset_right = page.width - paragraph['x1']

        if self.is_within_margin(offset_right, paragraph['x0'], 2):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.paragraph_format.left_indent = Pt(paragraph['x0'])

        if box is not None:
            box.append(p._element)
        return p

    def set_letter_spacing(self, run, spacing):
        print(spacing)
        r_element = run._element
        rPr = r_element.find(qn('w:rPr'))

        if rPr is None:
            rPr = OxmlElement('w:rPr')
            r_element.append(rPr)
        spacing_element = OxmlElement('w:spacing')
        spacing_element.set(qn('w:val'), str(spacing))
        rPr.append(spacing_element)

        return run

    def add_hyperlink(self, paragraph, run, hyperlink):
        if (hyperlink['uri']):
            part = paragraph.part
            r_id = part.relate_to(hyperlink['uri'], "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
            hyperlink = OxmlElement("w:hyperlink")
            hyperlink.set(qn("r:id"), r_id)
            hyperlink.append(run._element)
            paragraph._element.append(hyperlink)
        return run


    def rgb_color(self, color_data):
        rgb_color = tuple(int(c * 255) for c in color_data)
        if len(rgb_color) == 3:
            return RGBColor(*rgb_color)
        return None

    def add_rect(self, rect_data, page):
        crop = None
        p = self.get_empty_paragraph()
        run = p.add_run()
        shape_id = run.part.next_id
        x_cm, y_cm, width, height = self.get_pos_data(rect_data)
        fill_color = self.rgb_color(rect_data['non_stroking_color'])
        stroke_color = self.rgb_color(rect_data['stroking_color'])
        is_text_box = False
        valign = 't'
        text_box_content = None

        if 'text' in rect_data and len(rect_data['text']) > 0:
            bbox = [rect_data['x0'], rect_data['top'], rect_data['x1'], rect_data['bottom']]
            crop = page.crop(bbox)
            is_text_box = True
            bbox = {}
            for paragraph in rect_data['text']:
                if not 'top' in bbox:
                    bbox['top'] = paragraph['top']
                bbox['bottom'] = paragraph['bottom']
            if self.__parser.is_vertical_center(crop, bbox):
                valign = 'ctr'
        if rect_data['height'] > 2:
            anchor, text_box_content = CT_Rect.new(width, height, shape_id, x_cm, y_cm, fill_color, stroke_color, is_text_box, valign)
        else:
            anchor = CT_Line.new(width, height, shape_id, x_cm, y_cm, fill_color)

        if is_text_box:
            prev_paragraph = None
            for paragraph in rect_data['text']:
                set_offset = True
                if prev_paragraph is None and valign == 'ctr':
                    set_offset=False
                self.add_paragraph(paragraph, crop, prev_paragraph, text_box_content, set_offset)
                prev_paragraph = paragraph

        run._r.add_drawing(anchor)


    def process(self):
        with pdfplumber.open(self.__pdf_path, unicode_norm="NFC") as pdf:
            for page in pdf.pages:
                #if page.page_number !=2:
                    #continue
                print('page_number', page.page_number)
                self.__empty_p_count = 0
                self.__empty_p = None
                self.configure_page_from_pdf(page)
                page_data = self.__parser.get_data(page)
                prev_paragraph = None

                for rect in page_data['full_page_rects']:
                    self.add_rect(rect, page)

                for rect in page_data['rects']:
                    self.add_rect(rect, page)

                for image in page_data['images']:
                    img_path = self.extract_image(image)
                    self.add_image_block(image, img_path, page)

                for index, paragraph in enumerate(page_data['text']):
                    #print(paragraph)
                    self.add_paragraph(paragraph, page, prev_paragraph)

                    if prev_paragraph is None:
                        prev_paragraph = paragraph
                    elif prev_paragraph['top'] < paragraph['top']:
                        prev_paragraph = paragraph

                if page.page_number == 3:
                    break

        self.__doc.save(self.__output_path)
