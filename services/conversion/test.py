import base64
import os

import fitz  # PyMuPDF
import pdfplumber
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from PIL import Image
import io

def save_image_with_transparency(image_data, output_path, format):
    img = Image.open(io.BytesIO(image_data))
    rgba = img.convert("RGBA")
    datas = rgba.getdata()

    newData = []
    for item in datas:
        if item[0] == 0 and item[1] == 0 and item[2] == 0:  # finding black colour by its RGB value
            # storing a transparent value when we find a black colour
            newData.append((255, 255, 255, 0))
        else:
            newData.append(item)  # other colours remain unchanged

    rgba.putdata(newData)
    rgba.save(output_path, format)

def insert_image_to_doc(doc, image_path, bbox, xres, yres):
    """Вставляет изображение в Word с сохранением размеров и позиции."""
    x0, y0, x1, y1 = bbox
    pdf_width = x1 - x0
    pdf_height = y1 - y0

    width_in_inches = pdf_width / xres
    height_in_inches = pdf_height / yres

    left_margin = x0 / xres
    top_margin = y0 / yres

    # Добавляем пустые абзацы для имитации отступа сверху
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(top_margin * yres)  # Перевод в точки
    run = paragraph.add_run()

    run.add_picture(image_path, width=Inches(width_in_inches), height=Inches(height_in_inches))


def add_block_to_docx(doc, content, position, content_type="text", font_size=12):
    """
    Добавляет блок текста или изображения на определённые координаты в документе DOCX.

    :param doc: Объект документа.
    :param content: Текст или путь к изображению.
    :param position: Координаты блока (left, top, width, height) в EMU.
    :param content_type: Тип содержимого ("text" или "image").
    :param font_size: Размер шрифта для текста (Pt).
    """
    # Создаем XML-структуру для позиции
    left, top, width, height = position["left"], position["top"], position.get("width", 0), position.get("height", 0)
    drawing = parse_xml(
        f"""
        <w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                       distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="0" behindDoc="0"
                       locked="0" layoutInCell="1" allowOverlap="1">
                <wp:simplePos x="0" y="0"/>
                <wp:positionH relativeFrom="page">
                    <wp:posOffset>{left}</wp:posOffset>
                </wp:positionH>
                <wp:positionV relativeFrom="page">
                    <wp:posOffset>{top}</wp:posOffset>
                </wp:positionV>
                <wp:extent cx="{width}" cy="{height}"/>
                <wp:effectExtent l="0" t="0" r="0" b="0"/>
                <wp:wrapNone/>
                <wp:docPr id="1" name="Block"/>
                <wp:cNvGraphicFramePr/>
            </wp:anchor>
        </w:drawing>
        """
    )

    # Добавляем текст или изображение
    if content_type == "text":
        run = doc.add_paragraph().add_run(content)
        run.font.size = Pt(font_size)
        drawing.append(run._element)
    elif content_type == "image":
        run = doc.add_paragraph().add_run()
        image = run.add_picture(content, width=width, height=height)
        drawing.append(image._element)

    # Добавляем XML в документ
    doc.element.body.append(drawing)


def pdf_to_docx_coordinates(bbox, pdf_height, dpi=72):
    """
    Преобразует координаты из PDF в EMU для DOCX.

    :param bbox: Координаты блока из PDF [x0, y0, x1, y1] в точках.
    :param pdf_height: Высота страницы PDF в точках.
    :param dpi: Разрешение PDF (по умолчанию 72 dpi).
    :return: Словарь с координатами для DOCX в EMU.
    """
    x0, y0, x1, y1 = bbox
    # Переворот координат Y, так как PDF начинается с нижнего левого угла
    y0, y1 = pdf_height - y1, pdf_height - y0

    # Преобразуем точки в EMU (1 дюйм = 72 точки, 1 дюйм = 914400 EMU)
    emu_per_point = 914400 / dpi
    return {
        "left": int(x0 * emu_per_point),
        "top": int(y0 * emu_per_point),
        "width": int((x1 - x0) * emu_per_point),
        "height": int((y1 - y0) * emu_per_point),
    }

def configure_page_from_pdf(doc):
    """
    Настраивает параметры страницы Word на основе размеров PDF.
    :param doc: Объект документа Word
    :param pdf_path: Путь к PDF-файлу
    :param page_number: Номер страницы PDF (по умолчанию 0)
    """
    section = doc.sections[0]

    # Установка стандартных отступов (настраивайте по необходимости)
    section.left_margin = 0
    section.right_margin = 0
    section.top_margin = 0
    section.bottom_margin = 0



def insert_image_with_position(doc, image_path, bbox, pdf_page_height, xres=72, yres=72):
    """
    Вставляет изображение в Word с абсолютным позиционированием.

    :param doc: Объект документа DOCX.
    :param image_path: Путь к изображению.
    :param bbox: Координаты изображения в PDF (x0, y0, x1, y1).
    :param pdf_page_height: Высота страницы PDF в точках.
    :param xres: Горизонтальное разрешение изображения (DPI).
    :param yres: Вертикальное разрешение изображения (DPI).
    """
    if not os.path.exists(image_path):
        raise FileNotFoundError(f"Изображение не найдено: {image_path}")

    # Координаты и размеры из PDF
    x0, y0, x1, y1 = bbox
    pdf_width = x1 - x0
    pdf_height = y1 - y0

    # Переворот системы координат (PDF: нижний левый угол -> DOCX: верхний левый угол)
    y0 = pdf_page_height - y1

    # Перевод размеров из точек в EMU
    emu_per_point = 914400 / 72  # 1 точка = 1/72 дюйма, 1 дюйм = 914400 EMU
    left = int(x0 * emu_per_point / xres * 72)
    top = int(y0 * emu_per_point / yres * 72)
    width = int(pdf_width * emu_per_point / xres)
    height = int(pdf_height * emu_per_point / yres)

    # Добавляем изображение через add_picture для регистрации rId
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    inline_shape = run.add_picture(image_path, width=Inches(pdf_width / xres), height=Inches(pdf_height / yres))
    rId = inline_shape._inline.graphic.graphicData.pic.blipFill.blip.embed

    # Создаём XML для изображения с абсолютным позиционированием
    drawing = parse_xml(
        f"""
        <w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                       distT="0" distB="0" distL="0" distR="0"
                       simplePos="0" relativeHeight="0" behindDoc="0"
                       locked="0" layoutInCell="1" allowOverlap="1">
                <wp:simplePos x="0" y="0"/>
                <wp:positionH relativeFrom="page">
                    <wp:posOffset>{left}</wp:posOffset>
                </wp:positionH>
                <wp:positionV relativeFrom="page">
                    <wp:posOffset>{top}</wp:posOffset>
                </wp:positionV>
                <wp:extent cx="{width}" cy="{height}"/>
                <wp:effectExtent l="0" t="0" r="0" b="0"/>
                <wp:wrapNone/>
                <wp:docPr id="1" name="Image"/>
                <wp:cNvGraphicFramePr/>
                <a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
                    <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
                        <pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
                            <pic:nvPicPr>
                                <pic:cNvPr id="0" name="Picture"/>
                                <pic:cNvPicPr/>
                            </pic:nvPicPr>
                            <pic:blipFill>
                                <a:blip r:embed="{rId}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/>
                                <a:stretch>
                                    <a:fillRect/>
                                </a:stretch>
                            </pic:blipFill>
                            <pic:spPr>
                                <a:xfrm>
                                    <a:off x="0" y="0"/>
                                    <a:ext cx="{width}" cy="{height}"/>
                                </a:xfrm>
                                <a:prstGeom prst="rect">
                                    <a:avLst/>
                                </a:prstGeom>
                            </pic:spPr>
                        </pic:pic>
                    </a:graphicData>
                </a:graphic>
            </wp:anchor>
        </w:drawing>
        """
    )

    # Заменяем содержимое абзаца на XML с абсолютным позиционированием
    paragraph._element.clear_content()
    paragraph._element.append(drawing)


def get_pdf_page_size(pdf_path, page_number=1):
    """
    Определяет размеры страницы PDF в точках.

    :param pdf_path: Путь к PDF-файлу.
    :param page_number: Номер страницы (с 1).
    :return: Ширина и высота страницы в точках.
    """
    with pdfplumber.open(pdf_path) as pdf:
        page = pdf.pages[page_number - 1]  # Индексация страниц начинается с 0
        width = page.width  # Ширина страницы в точках
        height = page.height  # Высота страницы в точках
        return width, height

import fitz
from docx import Document
from docx.shared import Cm
from PIL import Image
import io

def configure_page_from_pdf(doc):
    """
    Настройка страниц Word-документа для соответствия размеру PDF.
    :param doc: Word-документ
    """
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

def get_pdf_page_size(pdf_path, page_number):
    """
    Возвращает размеры страницы PDF в точках.
    :param pdf_path: Путь к PDF-файлу
    :param page_number: Номер страницы (начиная с 1)
    :return: (width, height) в точках
    """
    doc = fitz.open(pdf_path)
    page = doc[page_number - 1]
    width, height = page.rect.width, page.rect.height
    doc.close()
    return width, height

folder_path = "tmp/s.pdf"
output_path = "tmp/s.zip"

pdf_document = fitz.open(folder_path)
doc = Document()
configure_page_from_pdf(doc)

page_num = 0
page = pdf_document[page_num]
width, height = get_pdf_page_size(folder_path, page_num + 1)

text_blocks = page.get_text("dict")["blocks"]
for block in text_blocks:
    if block['type'] == 1:  # Тип блока 1 означает изображение
        image_data = block['image']
        bbox = block['bbox']

        # Получаем изображение из байтов
        image_bytes = io.BytesIO(image_data)
        image = Image.open(image_bytes)

        # Рассчитываем размеры изображения в сантиметрах
        img_width = (bbox[2] - bbox[0]) * 0.0353  # Преобразуем из точек в см
        img_height = (bbox[3] - bbox[1]) * 0.0353

        # Вставляем изображение в Word
        run = doc.add_paragraph().add_run()
        run.add_picture(image_bytes, width=Cm(img_width), height=Cm(img_height))

print(text_blocks)

pdf_document.close()
doc.save(output_path)


